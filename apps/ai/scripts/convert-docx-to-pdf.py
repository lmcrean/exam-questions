import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Error: python-docx library not found.")
    print("Install it with: pip install python-docx")
    sys.exit(1)

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
except ImportError:
    print("Error: reportlab library not found.")
    print("Install it with: pip install reportlab")
    sys.exit(1)

def clean_text(text):
    """Clean text for PDF output."""
    if not text:
        return ""
    # Replace problematic characters
    replacements = {
        '\u2013': '-',
        '\u2014': '--',
        '\u2018': "'",
        '\u2019': "'",
        '\u201c': '"',
        '\u201d': '"',
        '\u2022': '-',
        '\u2026': '...',
        '&': '&amp;',
        '<': '&lt;',
        '>': '&gt;',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text

def convert_docx_to_pdf(docx_path):
    """Convert a Word document to PDF using reportlab."""
    docx_file = Path(docx_path)

    if not docx_file.exists():
        print(f"Error: File not found: {docx_path}")
        return False

    pdf_file = docx_file.with_suffix('.pdf')

    print(f"Converting: {docx_file.name}")
    print(f"To: {pdf_file.name}")
    print("Note: Using reportlab for conversion.")

    try:
        # Read the Word document
        doc = Document(docx_file)

        # Create PDF
        pdf = SimpleDocTemplate(
            str(pdf_file),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )

        # Container for PDF elements
        story = []

        # Get styles
        styles = getSampleStyleSheet()
        normal_style = styles['Normal']
        heading_style = styles['Heading1']

        # Process each paragraph
        for para in doc.paragraphs:
            text = clean_text(para.text.strip())

            if not text:
                story.append(Spacer(1, 0.2 * inch))
                continue

            # Determine style
            if para.style.name.startswith('Heading'):
                style = heading_style
            else:
                style = normal_style

            # Add paragraph
            p = Paragraph(text, style)
            story.append(p)
            story.append(Spacer(1, 0.1 * inch))

        # Process tables
        for table in doc.tables:
            table_data = []

            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = clean_text(cell.text.strip())
                    row_data.append(cell_text if cell_text else "")
                if any(row_data):  # Only add non-empty rows
                    table_data.append(row_data)

            if table_data:
                # Create table
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                ]))
                story.append(t)
                story.append(Spacer(1, 0.2 * inch))

        # Build PDF
        pdf.build(story)
        print(f"Successfully created: {pdf_file}")
        return True

    except Exception as e:
        print(f"Error during conversion: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    import sys

    # Allow command line argument or use default
    if len(sys.argv) > 1:
        docx_path = sys.argv[1]
    else:
        docx_path = r"C:\Projects\exam-questions\.notes\Student1 - Database Worksheet 1b Gotcha!.docx"

    convert_docx_to_pdf(docx_path)
